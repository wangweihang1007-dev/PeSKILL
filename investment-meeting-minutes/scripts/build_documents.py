from __future__ import annotations

import argparse
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SECTION_ORDER = [
    "1 公司定位", "2.1 产品", "2.2 市场情况", "2.3 核心客户",
    "3.1 核心技术体系", "3.2 技术差异化优势", "4 财务情况",
    "5.1 历史融资", "5.2 本轮融资安排", "6 发展计划",
]

SECTION_IMAGE_KEYWORDS = {
    "2.1 产品": ("产品", "业务", "解决方案", "应用", "产线", "设备"),
    "2.2 市场情况": ("市场", "行业", "规模", "增长", "竞争", "份额"),
    "2.3 核心客户": ("客户", "合作", "订单", "供应商", "产业链"),
    "3.1 核心技术体系": ("技术", "架构", "工艺", "原理", "研发", "专利"),
    "3.2 技术差异化优势": ("优势", "性能", "指标", "壁垒", "对比", "领先"),
    "4 财务情况": ("收入", "利润", "财务", "营收", "毛利", "预测"),
    "5.1 历史融资": ("历史融资", "股权", "股东", "融资历程"),
    "5.2 本轮融资安排": ("本轮融资", "融资计划", "估值", "募资", "资金用途"),
    "6 发展计划": ("规划", "计划", "里程碑", "产能", "未来", "发展"),
}
MAX_MINUTES_IMAGES = 6


def fresh(template: Path) -> Document:
    doc = Document(template)
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    return doc


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def select_section_images(images: list[dict]) -> dict[str, dict]:
    selected: dict[str, dict] = {}
    used: set[str] = set()
    for section, keywords in SECTION_IMAGE_KEYWORDS.items():
        candidates = []
        for item in images:
            if item.get("sha256") in used:
                continue
            context = item.get("context", "")
            score = sum(context.count(keyword) for keyword in keywords)
            if score:
                candidates.append((score, -int(item.get("source_page", 0)), item))
        if candidates and len(selected) < MAX_MINUTES_IMAGES:
            item = max(candidates, key=lambda candidate: (candidate[0], candidate[1]))[2]
            selected[section] = item
            used.add(item.get("sha256", ""))
    return selected


def add_section_image(doc: Document, item: dict) -> None:
    path = Path(item["file"])
    if not path.is_file():
        return
    try:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(path), width=Inches(5.8))
        caption = doc.add_paragraph(
            f"图：BP 第{item.get('source_page', '?')}页（{Path(item.get('source_file', '')).name}）"
        )
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.size = Pt(9)
    except Exception:
        return


def output_dir(project: Path) -> Path:
    base = project / "AI会议纪要输出"
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    candidate = base / stamp
    index = 2
    while candidate.exists():
        candidate = base / f"{stamp}-{index:02d}"
        index += 1
    candidate.mkdir(parents=True)
    return candidate


def build_corrected(template: Path, corrected: str, output: Path) -> None:
    doc = fresh(template); add_title(doc, "修正转录")
    for line in corrected.splitlines():
        if line.strip(): add_paragraph(doc, line.strip())
    doc.save(output)


def build_qa(template: Path, qa: dict, output: Path, title: str = "Q&A 整理") -> None:
    doc = fresh(template); add_title(doc, title)
    for item in qa["items"]:
        add_paragraph(doc, "Q：" + item["question"].strip(), "Q：")
        answer = "A：" + item["answer"].strip()
        if item.get("needs_verification") and "待核实" not in answer:
            answer += "（待核实）"
        add_paragraph(doc, answer, "A：")
    doc.save(output)


def build_minutes(template: Path, minutes: dict, qa: dict, output: Path, images: list[dict]) -> None:
    doc = fresh(template); add_title(doc, minutes.get("company_name") or minutes.get("project_name") or "项目会议纪要")
    section_images = select_section_images(images)
    participants = minutes.get("participants") or ["待确认"]
    metadata = [
        ("会议主题：", minutes.get("meeting_topic") or "待确认"),
        ("会议目的：", minutes.get("meeting_purpose") or "待确认"),
        ("记录时间：", minutes.get("meeting_date") or "待确认"),
        ("参会人：", "；".join(participants)),
    ]
    for label, value in metadata: add_paragraph(doc, label + value, label)
    for key in SECTION_ORDER:
        level = 2 if re.match(r"^[235]\.\d", key) else 1
        doc.add_paragraph(key, style=f"Heading {level}")
        content = minutes.get("sections", {}).get(key) or ["现有材料未提供相关信息，待确认。"]
        for paragraph in content: add_paragraph(doc, paragraph)
        if key in section_images:
            add_section_image(doc, section_images[key])
    doc.add_paragraph("7 访谈记录", style="Heading 1")
    for item in qa["items"]:
        add_paragraph(doc, "Q：" + item["question"].strip(), "Q：")
        answer = "A：" + item["answer"].strip()
        if item.get("needs_verification") and "待核实" not in answer: answer += "（待核实）"
        add_paragraph(doc, answer, "A：")
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--work-dir", required=True)
    parser.add_argument("--project", required=True)
    parser.add_argument("--template", required=True)
    args = parser.parse_args()
    work, project, template = Path(args.work_dir), Path(args.project), Path(args.template)
    corrected = (work / "corrected_transcript.txt").read_text(encoding="utf-8")
    qa = json.loads((work / "qa.json").read_text(encoding="utf-8"))
    minutes = json.loads((work / "minutes.json").read_text(encoding="utf-8"))
    images_path = work / "background_images.json"
    images = json.loads(images_path.read_text(encoding="utf-8")) if images_path.exists() else []
    out = output_dir(project)
    safe_name = re.sub(r'[<>:"/\\|?*]', "_", minutes.get("project_name") or "项目")
    build_corrected(template, corrected, out / "01_修正转录.docx")
    build_qa(template, qa, out / "02_QA整理.docx")
    build_minutes(template, minutes, qa, out / f"03_{safe_name}会议纪要.docx", images)
    print(out)


if __name__ == "__main__":
    main()
