from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from pypdf import PdfReader

from common import TRANSCRIPT_EXCLUDES, hidden_or_ignored, normalize_text, split_complete_turns


def read_docx(path: Path) -> str:
    doc = Document(path)
    blocks = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        blocks.extend(" | ".join(c.text.strip() for c in row.cells) for row in table.rows)
    return normalize_text("\n".join(blocks))


def read_pdf(path: Path) -> str:
    return normalize_text("\n".join(page.extract_text() or "" for page in PdfReader(path).pages))


def read_pptx(path: Path) -> str:
    slides: list[str] = []
    with zipfile.ZipFile(path) as archive:
        names = [n for n in archive.namelist() if re.fullmatch(r"ppt/slides/slide\d+\.xml", n)]
        names.sort(key=lambda n: int(re.search(r"\d+", Path(n).stem).group()))
        for index, name in enumerate(names, 1):
            root = ElementTree.fromstring(archive.read(name))
            text = [node.text.strip() for node in root.iter() if node.tag.endswith("}t") and node.text and node.text.strip()]
            slides.append(f"[第{index}页]\n" + "\n".join(text))
    return normalize_text("\n\n".join(slides))


def transcript_score(path: Path) -> int:
    name = path.stem.lower()
    priorities = (("初稿", 50), ("原文", 40), ("转录", 30), ("录音修正", 20), ("修正后结果", 10))
    return max((score for term, score in priorities if term in name), default=0)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--project", required=True)
    parser.add_argument("--work-dir", required=True)
    args = parser.parse_args()
    project = Path(args.project).resolve()
    work = Path(args.work_dir).resolve()
    work.mkdir(parents=True, exist_ok=True)
    manifest = {"status": "ok", "project": str(project), "reasons": [], "ignored": []}
    if not project.is_dir():
        manifest.update(status="blocked", reasons=["项目文件夹不存在或不可读"])
        return write_manifest(work, manifest, 2)

    files = [p for p in project.rglob("*") if p.is_file()]
    usable = [p for p in files if not hidden_or_ignored(p, project)]
    manifest["ignored"] = [str(p) for p in files if p not in usable]
    candidates = [
        p for p in usable
        if p.suffix.lower() in {".docx", ".txt"}
        and not any(term in p.stem.lower() for term in TRANSCRIPT_EXCLUDES)
    ]
    backgrounds = [p for p in usable if p.suffix.lower() in {".pdf", ".pptx"}]
    if not candidates:
        manifest["reasons"].append("未找到可用的 DOCX/TXT 转录初稿")
    if not backgrounds:
        manifest["reasons"].append("未找到可读的 PPTX/PDF 背景材料")
    selected = None
    if candidates:
        top = max(transcript_score(p) for p in candidates)
        top_items = [p for p in candidates if transcript_score(p) == top]
        if len(top_items) != 1:
            manifest["reasons"].append("存在多个无法唯一判断的转录初稿：" + "；".join(str(p) for p in top_items))
        else:
            selected = top_items[0]
    if manifest["reasons"]:
        manifest["status"] = "blocked"
        manifest["transcript_candidates"] = [str(p) for p in candidates]
        manifest["background_files"] = [str(p) for p in backgrounds]
        return write_manifest(work, manifest, 2)

    transcript = normalize_text(selected.read_text(encoding="utf-8-sig") if selected.suffix.lower() == ".txt" else read_docx(selected))
    background_parts: list[str] = []
    background_index: list[dict] = []
    for path in backgrounds:
        try:
            text = read_pdf(path) if path.suffix.lower() == ".pdf" else read_pptx(path)
        except Exception as exc:
            text = ""
            background_index.append({"file": str(path), "characters": 0, "error": str(exc)})
            continue
        background_index.append({"file": str(path), "characters": len(text)})
        if text:
            background_parts.append(f"===== {path.name} =====\n{text}")
    if not background_parts:
        manifest.update(status="blocked", reasons=["背景材料无法提取文字；请提供可读 PPTX/PDF，第一版不执行 OCR"])
        manifest["background_files"] = [str(p) for p in backgrounds]
        return write_manifest(work, manifest, 2)

    (work / "transcript.txt").write_text(transcript, encoding="utf-8")
    (work / "background.txt").write_text("\n\n".join(background_parts), encoding="utf-8")
    (work / "background_index.json").write_text(json.dumps(background_index, ensure_ascii=False, indent=2), encoding="utf-8")
    chunks = split_complete_turns(transcript)
    (work / "correction_chunks.json").write_text(json.dumps(chunks, ensure_ascii=False, indent=2), encoding="utf-8")
    manifest.update(transcript_file=str(selected), background_files=[str(p) for p in backgrounds], chunks=len(chunks))
    return write_manifest(work, manifest, 0)


def write_manifest(work: Path, manifest: dict, code: int) -> int:
    (work / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(manifest, ensure_ascii=False, indent=2))
    return code


if __name__ == "__main__":
    sys.exit(main())
